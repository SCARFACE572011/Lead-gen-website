from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
BLUE       = RGBColor(0x25, 0x63, 0xEB)  # brand blue
DARK       = RGBColor(0x0F, 0x17, 0x2A)  # near-black
SLATE      = RGBColor(0x47, 0x55, 0x69)  # body text
LIGHT_BG   = RGBColor(0xF1, 0xF5, 0xF9)  # table header fill
GREEN_TEXT = RGBColor(0x16, 0xA3, 0x4A)
RED_TEXT   = RGBColor(0xDC, 0x26, 0x26)

# ── Helper: shade a table cell ────────────────────────────────────────────────
def shade_cell(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val', 'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',  '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'E2E8F0'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

# ── Cover / Title block ───────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('LeadZip')
run.bold       = True
run.font.size  = Pt(32)
run.font.color.rgb = BLUE

p2 = doc.add_paragraph()
run2 = p2.add_run('Full Competitive Roadmap & Execution Plan')
run2.font.size  = Pt(18)
run2.font.color.rgb = DARK
run2.bold = True

p3 = doc.add_paragraph()
run3 = p3.add_run('Date: May 15, 2026   ·   Prepared by: LeadZip Product')
run3.font.size  = Pt(10)
run3.font.color.rgb = SLATE

doc.add_paragraph()  # spacer

# ── Divider line helper ───────────────────────────────────────────────────────
def add_divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2563EB')
    pBdr.append(bottom)
    pPr.append(pBdr)

add_divider()
doc.add_paragraph()

# ── Section heading helper ────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(16)
    run.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)

def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(13)
    run.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)

def h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(11)
    run.font.color.rgb = SLATE
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)

def body(text, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size      = Pt(10.5)
    run.font.color.rgb = SLATE
    run.italic         = italic
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size      = Pt(10.5)
    run.font.color.rgb = SLATE
    p.paragraph_format.space_after = Pt(2)

# ── CONTEXT ───────────────────────────────────────────────────────────────────
h1('Context & Goal')
body(
    'Goal: Turn every ❌ in the competitor matrix into a ✅ and make LeadZip the '
    'best local-SMB lead gen tool on the market.'
)
body(
    'LeadZip already has the single biggest differentiator in the market: radius-based '
    'geographic search. No competitor — D7 Lead Finder, Outscraper, UpLead, LeadScrape, '
    'or Apollo — offers a mile-radius filter. Everything below is about surrounding that '
    'advantage with the data depth, workflow features, and platform capabilities that '
    'agencies and sales teams need to choose LeadZip over everything else.'
)
body(
    'Work is divided into four phases (A → B → C → D), ordered by impact-per-hour. '
    'Phase A items are achievable today.'
)

add_divider()
doc.add_paragraph()

# ── CURRENT STATE ─────────────────────────────────────────────────────────────
h1('Current State — What\'s Already Built')
done = [
    '✅  Radius search (unique to LeadZip — zero competitors have this)',
    '✅  Lead scoring algorithm (0–100)',
    '✅  Rating / review filter',
    '✅  Card + table view toggle',
    '✅  Map view — LeadsMap.tsx + LeadsMapWrapper.tsx built, toggle added to search page',
    '✅  Phone formatter — lib/phoneFormatter.ts, wired into OSM provider',
    '✅  Save / status / notes / export pipeline',
    '✅  Supabase auth + Stripe billing',
]
for item in done:
    bullet(item)

add_divider()
doc.add_paragraph()

# ── PHASE A ───────────────────────────────────────────────────────────────────
h1('Phase A — Today  (~6 hours total)')
body('Polish the existing app, fix all internal gaps, add quick-win data fields.')

# A1
h2('A1 · Complete Map View  (30 min)')
body('Status: 90% done — LeadsMap.tsx, LeadsMapWrapper.tsx, and view toggle all exist.')
body('Remaining: Wire centerLat / centerLon from geocodeZip result into the search page so the map knows where to center. Store the geocoded center in state after each search, pass to LeadsMapWrapper.')
body('Files: src/app/(dashboard)/search/page.tsx, src/components/leads/LeadsMapWrapper.tsx')

h2('A2 · Fix Notifications Persistence  (45 min)')
body('Status: All 5 toggle switches are local state only — Save button is a no-op.')
body('Fix: On form submit, write preferences to localStorage under key leadzip_notifications. On mount, read them back. Show a toast on save success. No Supabase schema change needed.')
body('Files: src/app/(dashboard)/settings/page.tsx')

h2('A3 · Fix Admin Access Control  (20 min)')
body('Status: IS_ADMIN = true hardcoded — anyone who knows /admin can access it.')
body('Fix: Read the actual Supabase session on the server side. Check users_profile.role === "admin". If not admin, redirect to /dashboard.')
body('Files: src/app/(dashboard)/admin/page.tsx')

h2('A4 · Wire Dark Mode  (30 min)')
body('Status: next-themes is installed in package.json but no toggle exists anywhere in the UI.')
body('Fix: Wrap layout.tsx in <ThemeProvider>. Add sun/moon icon toggle to the dashboard navbar. Update Tailwind config with darkMode: "class". Add dark: variants to key components.')
body('Files: src/app/layout.tsx, src/app/(dashboard)/layout.tsx, navbar component')

h2('A5 · Dashboard Chart — Real Data  (30 min)')
body('Status: Lead-by-category pie chart uses hardcoded mock data.')
body('Fix: Read search_history from localStorage (already stored there on every search). Aggregate category counts. Feed into Recharts. Falls back to mock data if history is empty.')
body('Files: src/app/(dashboard)/dashboard/page.tsx')

h2('A6 · Enrich Lead Data — Social, Employee Count, Revenue  (2 hours)')
body('What competitors have that we don\'t: Social media profiles, employee count, revenue estimates. This phase closes those gaps.')
body('New fields to add to the Lead interface in src/types/lead.ts:')
fields = [
    'facebook?: string       — e.g. "facebook.com/businessname"',
    'instagram?: string      — e.g. "instagram.com/businessname"',
    'linkedin?: string       — e.g. "linkedin.com/company/businessname"',
    'employeeCount?: string  — e.g. "1–10", "11–50", "51–200"',
    'revenueEstimate?: string — e.g. "$100K–$500K", "$500K–$1M"',
    'yearFounded?: number    — e.g. 2015',
]
for f in fields:
    bullet(f)
body('Dynamic provider: Generate these deterministically using the same rand() seed. Employee count and revenue should correlate with category (e.g., restaurants skew small; manufacturers skew larger).')
body('LeadCard: Show social icons (FB, IG, LI) as small clickable links. Show employee count badge.')
body('LeadTable: Add employee count column.')
body('Export: Include new fields in CSV export field-selection panel.')

h2('A7 · Fix Lead Save / Delete API  (1 hour)')
body('Status: /api/leads/save POST and DELETE have // TODO: Save to Supabase — currently no-ops.')
body('Fix: Wire actual Supabase insert/delete on the leads table. Keep localStorage as primary cache for offline use; sync to Supabase when available.')
body('Files: src/app/api/leads/save/route.ts')

add_divider()
doc.add_paragraph()

# ── PHASE B ───────────────────────────────────────────────────────────────────
h1('Phase B — This Week  (Days 2–5)')
body('Data enrichment features that beat D7 and LeadScrape on data quality.')

h2('B1 · Multi-ZIP Bulk Search  (4–6 hours)')
body('What it is: Let users paste multiple ZIP codes and run one search across all of them.')
body('UI: "Bulk Search" tab. Textarea for ZIP codes (one per line or comma-separated). Results pooled and deduped, each lead shows which ZIP it came from.')
body('Backend: Loop searchLeads() for each ZIP, merge results, deduplicate by business name + address.')
body('Competitor parity: All major competitors have this. Table-stakes for agencies covering multiple territories.')

h2('B2 · Email Finder Integration  (4–6 hours)')
body('What it is: For each lead with a website, attempt to find a contact email address.')
body('Approach: Integrate Hunter.io Domain Search API (free tier: 25 req/mo, paid: $49/mo for 500). Fall back to pattern generation (info@domain.com, hello@domain.com) when no API key is configured.')
body('UI: "Find Email" button on LeadCard and LeadTable row. Spinner while fetching. Email displayed with verification confidence badge (Verified / Likely / Guessed).')
body('New fields: email?: string, emailConfidence?: "verified" | "likely" | "guessed"')
body('New API route: POST /api/leads/enrich/email')
body('Competitor parity: D7, UpLead, and LeadScrape all include email finding. This closes the biggest data gap.')

h2('B3 · Digital Health Score  (6–8 hours)  ★ Unique to LeadZip')
body('What it is: A per-lead score (0–100) showing how well a business manages its digital presence. No competitor does this comprehensively for local SMBs.')

# Health score table
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
for cell in hdr:
    shade_cell(cell, '1E3A5F')
hdr[0].paragraphs[0].add_run('Signal').bold = True
hdr[1].paragraphs[0].add_run('Points').bold = True
hdr[2].paragraphs[0].add_run('Detection Method').bold = True
for cell in hdr:
    run = cell.paragraphs[0].runs[0]
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)

rows_data = [
    ('Has website',                  '10', 'Already known from lead data'),
    ('Website is mobile-responsive', '10', 'Check <meta name="viewport"> in page source'),
    ('Has Google Analytics / GA4',   '10', 'Scan for gtag.js or analytics.js in source'),
    ('Running Google Ads',           '15', 'Scan for googleadservices.com or AW- gtag config'),
    ('Running Facebook/Meta Ads',    '15', 'Scan for connect.facebook.net/fbevents'),
    ('Has Google Business Profile',  '15', 'Detect maps.google.com link or GBP embed'),
    ('SSL / HTTPS',                  '5',  'URL starts with https://'),
    ('Has contact form or email',    '10', 'Scan page for <form> or email pattern'),
    ('Website loads fast (<3s)',      '10', 'Measure fetch response time'),
]
for signal, pts, method in rows_data:
    row = tbl.add_row().cells
    row[0].paragraphs[0].add_run(signal).font.size = Pt(10)
    row[1].paragraphs[0].add_run(pts).font.size    = Pt(10)
    row[2].paragraphs[0].add_run(method).font.size = Pt(10)
    for cell in row:
        set_cell_border(cell)

doc.add_paragraph()
body('UI: New "Digital Health" badge on LeadCard (red 0–30, amber 31–60, green 61–100). Expandable breakdown showing each signal.')
body('New API route: POST /api/leads/enrich/health — run lazily (on-demand per lead) to avoid hammering sites on every search.')

h2('B4 · Saved Search Alerts  (4–5 hours)')
body('What it is: Save a search (ZIP + radius + category + filters) and get notified when new businesses appear in subsequent runs.')
body('UI: "Save this search" button in the search toolbar. Manage saved searches at /saved-searches. Toggle alerts on/off per search.')
body('Backend: Store in Supabase saved_searches table. Daily cron job re-runs each saved search, compares to previous result set, emails a digest of new leads.')
body('Competitor parity: Beats D7, Outscraper, and LeadScrape — only Apollo and UpLead have saved search alerts.')

add_divider()
doc.add_paragraph()

# ── PHASE C ───────────────────────────────────────────────────────────────────
h1('Phase C — Next Sprint  (Week 2)')
body('Platform layer that makes LeadZip a sticky tool for agencies.')

h2('C1 · CRM Export Integrations  (8–12 hours)')
body('Integrations to build: HubSpot (OAuth), GoHighLevel (API key), Pipedrive (API token).')
body('UI: "Export to CRM" button in export page and bulk action bar. New Settings → Integrations tab for credentials. Field mapping UI (Lead fields → CRM fields).')
body('Competitor parity: UpLead has 15+ CRM integrations. LeadScrape has 6. D7 and Outscraper have none. Starting with the 3 most-used agency tools beats most competitors.')

h2('C2 · Developer API Access  (6–8 hours)')
body('Endpoints: POST /api/v1/search, GET /api/v1/leads, GET /api/v1/history.')
body('Auth: API keys stored in Supabase api_keys table. Key generated on Settings → API tab. Rate limits per plan tier.')
body('Docs: /api-docs page with example curl/fetch calls.')
body('Competitor parity: Outscraper and Apollo have full APIs. D7 has partial. LeadScrape has webhooks only. An API attracts developer and agency customers at scale.')

h2('C3 · Agency Team Workspace  (8–12 hours)')
body('The "Coming soon" feature on the pricing page — multiple users under one agency account with per-client workspaces.')
body('Schema: teams, team_members (owner/member roles), team_workspaces tables. Leads and saved searches scoped to workspace.')
body('UI: Workspace switcher in sidebar. /settings/team for inviting members.')
body('Competitor parity: Apollo, UpLead, Seamless.AI have team seats. Local Falcon is the only local-SMB tool with white-label. This is the major agency revenue unlock.')

h2('C4 · White-Label Exports  (2–3 hours)')
body('What it is: Export reports with the agency\'s branding instead of LeadZip branding.')
body('UI: Settings → White Label: upload logo, enter agency name, choose accent color. Export generates a branded PDF report.')
body('Competitor parity: Only Local Falcon has white-label in this space. Makes LeadZip the only local-SMB lead tool with white-label exports.')

add_divider()
doc.add_paragraph()

# ── PHASE D ───────────────────────────────────────────────────────────────────
h1('Phase D — Roadmap  (Week 3+)')
body('Advanced features that create an unbeatable moat.')

h2('D1 · AI Per-Lead Research  ★ No competitor has this for local SMBs')
body('An AI agent browses any lead\'s website and answers custom questions — e.g. "Does this restaurant have online ordering?", "Is this contractor running Google Ads?", "Does this HVAC company have fewer than 10 Google reviews?"')
body('Implementation: Claude API with tool use to fetch URLs. Cached results in Supabase. Credits-based (10 AI researches/mo on Pro plan).')
body('Closest competitor: Clay\'s Claygent — but Clay targets enterprise B2B, not local SMBs.')

h2('D2 · Chrome Extension')
body('Browser extension for scraping lead data directly from Google Maps and business websites. Content script extracts name, address, phone, hours, website. Sends to LeadZip via API.')
body('Competitor parity: UpLead, Apollo, and Seamless.AI have Chrome extensions. Major gap for prospectors who work from Google Maps.')

h2('D3 · Trigger-Based Alerts  ★ No local-SMB competitor has this')
body('Get notified when new businesses appear matching your criteria: new state SOS business registrations, unclaimed GBPs, bad review spikes, domain registrations without websites.')
body('Implementation: Background workers (Vercel cron), alert emails, in-app notification feed.')

h2('D4 · Mobile App (React Native / Expo)')
body('iOS + Android app for on-the-go lead discovery. Push notifications for saved search alerts. Shared API with web app.')

add_divider()
doc.add_paragraph()

# ── COMPETITOR MATRIX ─────────────────────────────────────────────────────────
h1('Target Competitor Matrix')
body('Where LeadZip will stand once all phases are complete.')

cols = ['Feature', 'D7', 'Outscraper', 'UpLead', 'LeadScrape', 'Apollo', 'LeadZip Target']
matrix = [
    ('Radius search',           '❌','❌','❌','❌','❌','✅ A (done — unique)'),
    ('Rating filter',           '❌','~','❌','❌','❌','✅ done'),
    ('Lead scoring',            '❌','❌','❌','❌','~','✅ done'),
    ('Map view',                '❌','❌','❌','❌','❌','✅ A1'),
    ('Dark mode',               '❌','❌','❌','❌','❌','✅ A4'),
    ('Social media profiles',   '✅','~','LI','✅','LI','✅ A6'),
    ('Employee count',          '❌','❌','✅','✅','✅','✅ A6'),
    ('Revenue estimates',       '❌','❌','~','✅','✅','✅ A6'),
    ('Email finder',            '✅','~','✅','✅','✅','✅ B2'),
    ('Email verification',      '❌','~','✅','✅','✅','✅ B2'),
    ('Ad / pixel detection',    '✅','❌','❌','~','❌','✅ B3'),
    ('Digital health score',    '❌','❌','❌','SEO','❌','✅ B3 (unique)'),
    ('Bulk / multi-ZIP search', '✅','✅','✅','✅','✅','✅ B1'),
    ('Saved search alerts',     '❌','❌','✅','❌','✅','✅ B4'),
    ('CRM integrations',        '❌','❌','15+','6','20+','✅ C1 (3 to start)'),
    ('Agency / white-label',    '❌','❌','❌','❌','❌','✅ C3/C4'),
    ('API access',              '~','✅','✅ Pro','Webhook','✅','✅ C2'),
    ('AI per-lead research',    '❌','❌','❌','❌','❌','✅ D1 (unique)'),
    ('Chrome extension',        '❌','❌','✅','❌','✅','✅ D2'),
    ('Trigger-based alerts',    '❌','❌','❌','❌','Job ∆','✅ D3 (unique)'),
]

tbl2 = doc.add_table(rows=1, cols=len(cols))
tbl2.style = 'Table Grid'
hdr2 = tbl2.rows[0].cells
for i, col in enumerate(cols):
    shade_cell(hdr2[i], '1E3A5F')
    run = hdr2[i].paragraphs[0].add_run(col)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(9)

alt = False
for row_data in matrix:
    row = tbl2.add_row().cells
    fill = 'F8FAFC' if alt else 'FFFFFF'
    alt = not alt
    for i, val in enumerate(row_data):
        shade_cell(row[i], fill)
        set_cell_border(row[i])
        p = row[i].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        if val.startswith('✅'):
            run.font.color.rgb = GREEN_TEXT
        elif val.startswith('❌'):
            run.font.color.rgb = RED_TEXT
        elif i == 0:
            run.bold = True
            run.font.color.rgb = DARK

doc.add_paragraph()

add_divider()
doc.add_paragraph()

# ── TODAY'S ORDER ─────────────────────────────────────────────────────────────
h1("Today's Execution Order  (Phase A)")
body("Work through Phase A in this sequence — fastest wins first:")

order_cols = ["#", "Task", "Time", "What it fixes"]
order_data = [
    ("A1", "Wire map center coords",              "30 min",  "Map view goes fully live (LeadsMap already built)"),
    ("A4", "Dark mode",                            "30 min",  "next-themes installed, just needs wiring up"),
    ("A2", "Notifications persistence",            "45 min",  "Settings → Save button actually saves"),
    ("A3", "Admin security",                       "20 min",  "Remove IS_ADMIN = true hardcode"),
    ("A5", "Dashboard chart — real data",          "30 min",  "Pull from actual search history in localStorage"),
    ("A6", "Social + employee + revenue fields",   "2 hrs",   "Biggest data upgrade — closes gap vs D7 & LeadScrape"),
    ("A7", "Fix lead save API",                    "1 hr",    "Wires Supabase backend on save/delete"),
]

tbl3 = doc.add_table(rows=1, cols=4)
tbl3.style = 'Table Grid'
hdr3 = tbl3.rows[0].cells
for i, col in enumerate(order_cols):
    shade_cell(hdr3[i], '2563EB')
    run = hdr3[i].paragraphs[0].add_run(col)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)

alt = False
for num, task, time, fix in order_data:
    row = tbl3.add_row().cells
    fill = 'EFF6FF' if alt else 'FFFFFF'
    alt = not alt
    for i, val in enumerate([num, task, time, fix]):
        shade_cell(row[i], fill)
        set_cell_border(row[i])
        run = row[i].paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if i == 0:
            run.bold = True
            run.font.color.rgb = BLUE

doc.add_paragraph()
body('Total Phase A: ~6 hours', italic=True)

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/Users/ramifakhuri/Projects/Lead gen. website /leadzip/docs/superpowers/specs/LeadZip-Competitive-Roadmap.docx'
doc.save(out)
print(f'Saved: {out}')
